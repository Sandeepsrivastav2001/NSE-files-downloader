from docx import Document
from docx.enum.text import WD_COLOR_INDEX

# Input & Output file path
input_file = r"C:\Users\Administrator\Desktop\My.docx"       # Apni Word file ka naam
output_file = r"C:\Users\Administrator\Desktop\output_highlighted.docx"

# Box characters to detect (add more if needed)
box_chars = ["□", "■", "☐", "❑", "❏"]

# Load the Word file
doc = Document(input_file)

# Function to highlight runs with box characters
def highlight_boxes_in_runs(runs):
    for run in runs:
        for box in box_chars:
            if box in run.text:
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW

# Process paragraphs
for para in doc.paragraphs:
    highlight_boxes_in_runs(para.runs)

# Process tables (Word me boxes table ke cells me bhi ho sakte hain)
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                highlight_boxes_in_runs(para.runs)

# Save the updated document
doc.save(output_file)

print(f"Done! Highlighted file saved as: {output_file}")
